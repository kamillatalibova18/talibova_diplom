import os
import json
import datetime
import MySQLdb
from PyQt6.QtWidgets import (QMainWindow, QMessageBox, QListWidgetItem,
                             QInputDialog, QFileDialog, QCheckBox)
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from clerkkk import Clerk_Form
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


class Clerk_Window(QMainWindow):
    def __init__(self):
        super().__init__()
        try:
            self.conn = MySQLdb.connect("localhost", "root", "root", "gor_stroy", charset="utf8mb4")
            self.cursor = self.conn.cursor()
        except MySQLdb.Error as e:
            QMessageBox.critical(self, "Ошибка базы данных", f"Не удалось подключиться к базе данных: {e}")
            raise

        self.ui = Clerk_Form()
        self.ui.setupUi(self)

        # Загрузка данных
        self.load_projects()
        self.current_project_id = None

        # Подключение сигналов
        self.ui.listWidget_project.itemSelectionChanged.connect(self.on_project_selected)
        self.ui.pushButton_create_doc.clicked.connect(self.create_document)
        self.ui.pushButton_open_project.clicked.connect(self.open_project_folder)
        self.ui.lineEdit_search_document.textChanged.connect(self.filter_documents)
        self.ui.pushButton_preview.clicked.connect(self.preview_document)
        self.ui.pushButton_delete_doc.clicked.connect(self.delete_document)
        self.ui.pushButton_export.clicked.connect(self.export_documents)
        self.ui.pushButton_archive_project.clicked.connect(self.archive_project)
        self.ui.comboBox_filter_project.currentIndexChanged.connect(self.filter_projects)
        self.ui.lineEdit_search_project.textChanged.connect(self.filter_projects)
        self.ui.pushButton_refresh.clicked.connect(self.load_projects)

    def archive_project(self):
        """Архивирование проекта с расширенной логикой"""
        if not self.current_project_id:
            QMessageBox.warning(self, "Ошибка", "Сначала выберите проект")
            return

        try:
            # 1. Проверка, не архивный ли уже проект
            self.cursor.execute("SELECT id FROM Project_Archive WHERE project_id = %s",
                                (self.current_project_id,))
            if self.cursor.fetchone():
                QMessageBox.information(self, "Информация", "Этот проект уже в архиве")
                return

            # 2. Получение информации о проекте
            self.cursor.execute("SELECT name, status FROM Projects WHERE id = %s",
                                (self.current_project_id,))
            project_name, current_status = self.cursor.fetchone()

            # 3. Запрос подтверждения с дополнительными параметрами
            confirm_dialog = QMessageBox(self)
            confirm_dialog.setWindowTitle("Подтверждение архивации")
            confirm_dialog.setText(
                f"Вы уверены, что хотите архивировать проект '{project_name}'?\n"
                f"Текущий статус: {current_status}"
            )
            confirm_dialog.setStandardButtons(
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            confirm_dialog.setDefaultButton(QMessageBox.StandardButton.No)

            # Добавляем чекбокс для переноса документов
            transfer_checkbox = QCheckBox("Перенести документы в архивную папку")
            transfer_checkbox.setChecked(True)
            layout = confirm_dialog.layout()
            layout.addWidget(transfer_checkbox, 1, 0, 1, 2)

            if confirm_dialog.exec() != QMessageBox.StandardButton.Yes:
                return

            # 4. Запрос причины архивации
            reason, ok = QInputDialog.getText(
                self,
                "Причина архивации",
                "Укажите причину архивации проекта:",
                text="Проект завершен"
            )
            if not ok:
                return

            # 5. Начало транзакции
            self.conn.autocommit(False)

            try:
                # 6. Обновление статуса проекта
                self.cursor.execute("""
                    UPDATE Projects 
                    SET status = 'архивный', updated_at = %s 
                    WHERE id = %s
                """, (datetime.date.today(), self.current_project_id))

                # 7. Добавление записи в архив с причиной
                self.cursor.execute("""
                    INSERT INTO Project_Archive (project_id, archived_at, reason)
                    VALUES (%s, %s, %s)
                """, (self.current_project_id, datetime.date.today(), reason))



                # 9. Перенос документов (если выбран)
                if transfer_checkbox.isChecked():
                    self._transfer_project_documents()

                # 10. Фиксация изменений
                self.conn.commit()

                # 11. Обновление интерфейса и уведомление
                self.load_projects()
                QMessageBox.information(
                    self,
                    "Успех",
                    f"Проект '{project_name}' успешно архивирован\n"
                    f"Причина: {reason}"
                )

            except Exception as e:
                self.conn.rollback()
                raise e
            finally:
                self.conn.autocommit(True)

        except Exception as e:
            QMessageBox.critical(
                self,
                "Ошибка архивации",
                f"Не удалось архивировать проект:\n{str(e)}"
            )

    def _transfer_project_documents(self):
        """Внутренний метод для переноса документов проекта"""
        # 1. Создаем архивную папку
        archive_folder = os.path.join("archive", f"project_{self.current_project_id}")
        os.makedirs(archive_folder, exist_ok=True)

        # 2. Получаем все документы проекта
        self.cursor.execute("""
            SELECT id, file_path FROM Documents 
            WHERE project_id = %s
        """, (self.current_project_id,))
        documents = self.cursor.fetchall()

        # 3. Переносим каждый документ
        for doc_id, file_path in documents:
            if not file_path or not os.path.exists(file_path):
                continue
                # Формируем новое имя файла
                filename = os.path.basename(file_path)
                new_path = os.path.join(archive_folder, filename)

                # Уникальное имя, если файл уже существует
                counter = 1
                while os.path.exists(new_path):
                    name, ext = os.path.splitext(filename)
                    new_path = os.path.join(archive_folder, f"{name}_{counter}{ext}")
                    counter += 1

                # Переносим файл
                os.rename(file_path, new_path)

                # Обновляем путь в базе данных
                self.cursor.execute("""
                    UPDATE Documents 
                    SET file_path = %s 
                    WHERE id = %s
                """, (new_path, doc_id))

    def load_projects(self):
        """Загрузка всех проектов из базы данных"""
        try:
            self.ui.listWidget_project.clear()
            self.cursor.execute("SELECT id, name, status FROM Projects")

            self.all_projects = []  # Сохраняем все проекты для фильтрации
            for project_id, name, status in self.cursor.fetchall():
                self.all_projects.append({
                    'id': project_id,
                    'name': name,
                    'status': status
                })

            self.filter_projects()  # Применяем текущие фильтры

        except MySQLdb.Error as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить проекты: {e}")

    def filter_projects(self):
        """Фильтрация проектов по статусу и поисковому запросу"""
        if not hasattr(self, 'all_projects'):
            return

        search_text = self.ui.lineEdit_search_project.text().lower()
        status_filter = self.ui.comboBox_filter_project.currentText()

        self.ui.listWidget_project.clear()

        for project in self.all_projects:
            # Фильтрация по статусу
            status_match = True
            if status_filter == "Активные":
                status_match = project['status'] != 'архивный'
            elif status_filter == "Архивные":
                status_match = project['status'] == 'архивный'

            # Фильтрация по тексту поиска
            search_match = search_text in project['name'].lower()

            if status_match and search_match:
                item = QListWidgetItem(project['name'])
                item.setData(100, project['id'])  # Сохраняем ID проекта
                self.ui.listWidget_project.addItem(item)

    def on_project_selected(self):
        """Обработка выбора проекта"""
        item = self.ui.listWidget_project.currentItem()
        if not item:
            return

        project_id = item.data(100)
        self.current_project_id = project_id
        project_name = item.text()

        # Загрузка информации о проекте
        try:
            self.cursor.execute("""
                SELECT status, updated_at 
                FROM Projects 
                WHERE id = %s
            """, (project_id,))
            status, updated_at = self.cursor.fetchone()

            self.ui.label_status_project.setText(f"Статус: {status}")
            self.ui.label_data_change_project.setText(f"Обновлено: {updated_at.strftime('%d.%m.%Y')}")
            self.ui.textBrowser_preview.setText(f"Выбран проект: {project_name}\nСтатус: {status}")

            # Загрузка документов проекта
            self.load_project_documents(project_id)
        except MySQLdb.Error as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить данные проекта: {e}")

    def load_project_documents(self, project_id):
        """Загрузка документов выбранного проекта"""
        try:
            self.ui.listWidget_document.clear()
            self.cursor.execute("""
                SELECT d.id, dt.name, d.file_path, d.created_at 
                FROM Documents d
                JOIN Documents_Type dt ON d.type_id = dt.id
                WHERE d.project_id = %s
            """, (project_id,))

            for doc_id, doc_type, file_path, created_at in self.cursor.fetchall():
                item = QListWidgetItem(f"{doc_type} ({created_at.strftime('%d.%m.%Y')})")
                item.setData(100, doc_id)  # ID документа
                item.setData(101, file_path)  # Путь к файлу
                self.ui.listWidget_document.addItem(item)
        except MySQLdb.Error as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить документы: {e}")

    def filter_documents(self):
        """Фильтрация документов по тексту поиска"""
        search_text = self.ui.lineEdit_search_document.text().lower()

        for i in range(self.ui.listWidget_document.count()):
            item = self.ui.listWidget_document.item(i)
            item_text = item.text().lower()
            item.setHidden(search_text not in item_text)

    def open_project_folder(self):
        """Открытие папки проекта"""
        if not self.current_project_id:
            QMessageBox.warning(self, "Ошибка", "Сначала выберите проект")
            return

        try:
            # Получаем путь к чертежу проекта (как пример)
            self.cursor.execute("""
                SELECT d.file_path 
                FROM Projects p
                JOIN Drawings d ON p.id_drawing = d.id
                WHERE p.id = %s
            """, (self.current_project_id,))

            result = self.cursor.fetchone()
            if result and result[0]:
                file_path = result[0]
                folder_path = os.path.dirname(file_path)

                if os.path.exists(folder_path):
                    os.startfile(folder_path)
                else:
                    QMessageBox.warning(self, "Ошибка", "Папка проекта не найдена")
            else:
                QMessageBox.information(self, "Информация", "Для проекта не указана папка с файлами")
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось открыть папку проекта: {e}")

    def preview_document(self):
        """Предпросмотр выбранного документа"""
        item = self.ui.listWidget_document.currentItem()
        if not item:
            QMessageBox.warning(self, "Ошибка", "Выберите документ для предпросмотра")
            return

        file_path = item.data(101)
        if not file_path or not os.path.exists(file_path):
            QMessageBox.warning(self, "Ошибка", "Файл документа не найден")
            return

        try:
            if file_path.endswith('.docx'):
                doc = Document(file_path)
                preview_text = ""
                for para in doc.paragraphs:
                    preview_text += para.text + "\n"
                self.ui.textBrowser_preview.setPlainText(preview_text)
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    self.ui.textBrowser_preview.setPlainText(f.read())
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось прочитать документ: {e}")

    def delete_document(self):
        """Удаление выбранного документа"""
        item = self.ui.listWidget_document.currentItem()
        if not item:
            QMessageBox.warning(self, "Ошибка", "Выберите документ для удаления")
            return

        doc_id = item.data(100)
        file_path = item.data(101)

        reply = QMessageBox.question(
            self, 'Подтверждение',
            'Вы уверены, что хотите удалить этот документ?',
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            try:
                # Удаляем запись из БД
                self.cursor.execute("DELETE FROM Documents WHERE id = %s", (doc_id,))
                self.conn.commit()

                # Удаляем файл, если он существует
                if file_path and os.path.exists(file_path):
                    os.remove(file_path)

                # Обновляем список документов
                self.load_project_documents(self.current_project_id)
                QMessageBox.information(self, "Успех", "Документ успешно удален")
            except Exception as e:
                QMessageBox.warning(self, "Ошибка", f"Не удалось удалить документ: {e}")

    def export_documents(self):
        """Экспорт документов проекта в выбранную папку"""
        if not self.current_project_id:
            QMessageBox.warning(self, "Ошибка", "Сначала выберите проект")
            return

        # 1. Выбор папки для экспорта
        folder_path = QFileDialog.getExistingDirectory(
            self,
            "Выберите папку для экспорта",
            os.path.expanduser("~")  # Начинать с домашней директории пользователя
        )

        if not folder_path:  # Если пользователь отменил выбор
            return

        try:
            # 2. Получаем все документы текущего проекта из БД
            self.cursor.execute("""
                SELECT file_path FROM Documents 
                WHERE project_id = %s
            """, (self.current_project_id,))

            exported_files = 0
            errors = []

            # 3. Копируем каждый документ
            for (file_path,) in self.cursor.fetchall():
                if file_path and os.path.exists(file_path):
                    try:
                        file_name = os.path.basename(file_path)
                        dest_path = os.path.join(folder_path, file_name)

                        # Создаем уникальное имя файла, если такой уже существует
                        counter = 1
                        while os.path.exists(dest_path):
                            name, ext = os.path.splitext(file_name)
                            dest_path = os.path.join(folder_path, f"{name}_{counter}{ext}")
                            counter += 1

                        # Копируем файл
                        with open(file_path, 'rb') as src, open(dest_path, 'wb') as dst:
                            dst.write(src.read())
                        exported_files += 1
                    except Exception as e:
                        errors.append(f"Ошибка при копировании {file_path}: {str(e)}")

            # 4. Показываем результат
            message = f"Успешно экспортировано {exported_files} документов"
            if errors:
                message += "\n\nОшибки:\n" + "\n".join(errors)

            QMessageBox.information(
                self,
                "Экспорт завершен",
                message
            )

            # 5. Открываем папку с экспортированными документами
            if exported_files > 0:
                os.startfile(folder_path)

        except Exception as e:
            QMessageBox.warning(
                self,
                "Ошибка",
                f"Не удалось экспортировать документы: {e}"
            )

    def ensure_project_has_template(self, project_id):
        """Проверяет и создает запись в Documents с шаблоном по умолчанию при необходимости"""
        try:
            # Проверяем, есть ли уже документ у проекта
            self.cursor.execute("""
                SELECT id_tempalte FROM Documents 
                WHERE project_id = %s 
                LIMIT 1
            """, (project_id,))
            result = self.cursor.fetchone()

            if result:
                return result[0]  # Возвращаем ID существующего шаблона

            # Если документа нет, создаем его с шаблоном по умолчанию
            self.cursor.execute("""
                INSERT INTO Documents (project_id, type_id, file_path, created_at, id_tempalte)
                VALUES (%s, 1, '', %s, 
                (SELECT id FROM Templates WHERE name = 'частный дом' LIMIT 1))
            """, (project_id, datetime.date.today()))
            self.conn.commit()

            # Возвращаем ID использованного шаблона
            self.cursor.execute("SELECT id_tempalte FROM Documents WHERE project_id = %s", (project_id,))
            return self.cursor.fetchone()[0]

        except MySQLdb.Error as e:
            QMessageBox.warning(self, "Ошибка БД", f"Не удалось назначить шаблон: {e}")
            return None

    def create_document(self):
        item = self.ui.listWidget_project.currentItem()
        if not item:
            QMessageBox.warning(self, "Ошибка", "Выберите проект.")
            return

        project_name = item.text()

        doc_type, ok = QInputDialog.getItem(self, "Выбор документа", "Выберите тип документа:",
                                            ["Отчёт", "Договор"], 0, False)
        if not ok:
            return

        try:
            # Получаем основные данные проекта
            self.cursor.execute("""
                SELECT p.id, p.name, p.status, p.updated_at, p.id_drawing, p.id_model
                FROM Projects p
                WHERE p.name = %s
                LIMIT 1
            """, (project_name,))
            project = self.cursor.fetchone()

            if not project:
                QMessageBox.warning(self, "Ошибка", "Проект не найден в базе.")
                return

            project_id, name, status, updated_at, drawing_id, model_id = project

            # Убеждаемся, что у проекта есть шаблон (создаем запись в Documents при необходимости)
            template_id = self.ensure_project_has_template(project_id)
            if not template_id:
                QMessageBox.warning(self, "Ошибка", "Не удалось назначить шаблон проекту.")
                return

            # Для договоров шаблон не обязателен
            if doc_type == "Отчёт" and not template_id:
                QMessageBox.warning(self, "Ошибка", "Для создания отчета требуется шаблон.")
                return

            completion_date = updated_at.strftime('%d.%m.%Y') if updated_at else "не указана"
            engineer = "Иванов И.И."

            # Получаем параметры шаблона (только для отчетов)
            params = {}
            if doc_type == "Отчёт" and template_id:
                self.cursor.execute("SELECT json_data FROM Templates WHERE id = %s", (template_id,))
                row = self.cursor.fetchone()
                if row and row[0]:
                    try:
                        params = json.loads(row[0])
                    except json.JSONDecodeError:
                        QMessageBox.warning(self, "Ошибка", "Неверный формат данных шаблона.")
                        return

            # Получаем пути к чертежам и моделям
            drawing_path = ""
            if drawing_id:
                self.cursor.execute("SELECT file_path FROM Drawings WHERE id = %s", (drawing_id,))
                drawing_row = self.cursor.fetchone()
                drawing_path = drawing_row[0] if drawing_row else ""

            model_path = ""
            if model_id:
                self.cursor.execute("SELECT file_path FROM 3d_model WHERE id = %s", (model_id,))
                model_row = self.cursor.fetchone()
                model_path = model_row[0] if model_row else ""

            os.makedirs("documents", exist_ok=True)
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')

            if doc_type == "Отчёт":
                filename = f"documents/Отчет_{project_name}_{timestamp}.docx"
                doc = Document()
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)

                doc.add_paragraph("ООО \"Гор-Строй\"", style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                doc.add_heading("Отчёт о завершении строительного проекта", level=1)

                doc.add_paragraph(f"Наименование проекта: {project_name}")
                doc.add_paragraph(f"Статус: {status}")
                doc.add_paragraph(f"Дата завершения: {completion_date}")
                doc.add_paragraph(f"Ответственный инженер: {engineer}")

                if template_id:
                    self.cursor.execute("SELECT name FROM Templates WHERE id = %s", (template_id,))
                    template_name = self.cursor.fetchone()
                    doc.add_paragraph(f"Использованный шаблон: {template_name[0] if template_name else 'N/A'}")

                doc.add_heading("Технические характеристики", level=2)
                for k, v in params.items():
                    doc.add_paragraph(f"{k.replace('_', ' ').capitalize()}: {v}")

                doc.add_heading("Приложение", level=2)
                if drawing_path and os.path.exists(drawing_path):
                    try:
                        doc.add_paragraph("Чертёж:")
                        doc.add_picture(drawing_path, width=Inches(5.5))
                    except Exception as e:
                        doc.add_paragraph(f"Ошибка при добавлении чертежа: {str(e)}")
                doc.add_paragraph(f"3D-модель: {model_path if model_path else 'не указана'}")

                try:
                    doc.save(filename)
                    # Обновляем запись о документе в БД
                    doc_type_id = 1  # ID для типа "Отчет"
                    self.cursor.execute("""
                        UPDATE Documents 
                        SET type_id = %s, file_path = %s, created_at = %s 
                        WHERE project_id = %s
                    """, (doc_type_id, filename, datetime.date.today(), project_id))
                    self.conn.commit()
                    QMessageBox.information(self, "Успех", f"Документ создан:\n{filename}")
                except Exception as e:
                    QMessageBox.warning(self, "Ошибка", f"Не удалось сохранить документ: {e}")

            elif doc_type == "Договор":
                filename = f"documents/Договор_{project_name}_{timestamp}.docx"
                doc = Document()
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)

                # Стандартный текст договора
                doc.add_paragraph("ДОГОВОР № 17-СМР").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                doc.add_paragraph("на выполнение строительно-монтажных работ").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                doc.add_paragraph(f"г. Москва {' ' * 60} {datetime.date.today().strftime('%d.%m.%Y')}")
                doc.add_paragraph()

                doc.add_paragraph(
                    "Общество с ограниченной ответственностью \"ГорСтройПроект\", в лице Генерального директора Иванова И.И., "
                    "действующего на основании Устава, именуемое в дальнейшем \"Подрядчик\", с одной стороны, и Заказчик, "
                    "именуемый в дальнейшем \"Заказчик\", с другой стороны, совместно именуемые \"Стороны\", заключили "
                    "настоящий договор о нижеследующем:"
                )

                doc.add_heading("1. Предмет договора", level=1)
                doc.add_paragraph(
                    f"1.1. Подрядчик обязуется выполнить по заданию Заказчика строительно-монтажные работы в рамках проекта "
                    f"«{project_name}» в соответствии с условиями настоящего договора и передать результат работ Заказчику."
                )

                doc.add_heading("2. Сроки выполнения работ", level=1)
                doc.add_paragraph(
                    "2.1. Срок начала работ: с даты подписания договора.\n"
                    f"2.2. Срок окончания работ: {completion_date}."
                )

                doc.add_heading("3. Стоимость и порядок оплаты", level=1)
                doc.add_paragraph(
                    "3.1. Общая стоимость работ составляет 1 200 000 (Один миллион двести тысяч) рублей.\n"
                    "3.2. Оплата производится Заказчиком поэтапно согласно графику выполнения работ."
                )

                doc.add_heading("4. Ответственность сторон", level=1)
                doc.add_paragraph(
                    "4.1. За невыполнение условий договора стороны несут ответственность в соответствии с действующим "
                    "законодательством РФ."
                )

                doc.add_heading("5. Заключительные положения", level=1)
                doc.add_paragraph(
                    "5.1. Настоящий договор вступает в силу с момента его подписания сторонами.\n"
                    "5.2. Настоящий договор составлен в двух экземплярах, по одному для каждой из сторон."
                )

                doc.add_paragraph("\nПодписи сторон:\n")
                table = doc.add_table(rows=2, cols=2)
                table.cell(0, 0).text = "ПОДРЯДЧИК:\nООО «ГорСтройПроект»\nГенеральный директор: ____________"
                table.cell(0, 1).text = "ЗАКАЗЧИК:\n__________________________\nФИО: ____________"

                try:
                    doc.save(filename)
                    # Обновляем запись о документе в БД
                    doc_type_id = 2  # ID для типа "Договор"
                    self.cursor.execute("""
                        UPDATE Documents 
                        SET type_id = %s, file_path = %s, created_at = %s 
                        WHERE project_id = %s
                    """, (doc_type_id, filename, datetime.date.today(), project_id))
                    self.conn.commit()
                    QMessageBox.information(self, "Успех", f"Документ создан:\n{filename}")
                except Exception as e:
                    QMessageBox.warning(self, "Ошибка", f"Не удалось сохранить документ: {e}")

        except MySQLdb.Error as e:
            QMessageBox.warning(self, "Ошибка базы данных", f"Ошибка при работе с базой данных: {e}")
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Произошла непредвиденная ошибка: {e}")

